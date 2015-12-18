####################################################
#                  Revision: 1.0                   #
#              Updated on: 11/12/2015              #
####################################################

####################################################
#                                                  #
#   This script contains a functions usefull       #
#   to generate Raid Fault report.                 # 
#                                                  #
#   Author: Zankar Sanghavi                        #
#                                                  #
#   © Dot Hill Systems Corporation                 #
#                                                  #
####################################################


class Raid_Fault_Functions:
    
    ##############################
    # Find lines with given strings
    ##############################
    def extract_create_recreation_data (c_path,data, 
                                        creation_start_string, 
                                        creation_complete_string,
                                        recreation_start_string,
                                        recreation_complete_string):
        
        import sys
        sys.path.insert(0, r''+str(c_path)+'/IO Stress')
        import log_functions
        lf=log_functions.Log_Functions
        
        import numpy as np
        data = np.array(data)
        
        # search & extract start string data, creation 
        c_start_indices = lf.find_index(data, creation_start_string)
        creation_start = data[c_start_indices]
        

        # search & extract complete string data, creation completed 
        c_complete_indices = lf.find_index(data, creation_complete_string)
        creation_complete = data[c_complete_indices]

        
        # search & extract start string data, reconstruction 
        r_start_indices = lf.find_index(data, recreation_start_string)
        recreation_start = data[r_start_indices]
        

        # search & extract complete string data, reconstruction completed 
        r_complete_indices = lf.find_index(data, recreation_complete_string)
        recreation_complete = data[r_complete_indices]

        return [creation_start, creation_complete, 
                recreation_start, recreation_complete]

                
     
    ##########################
    # This function will take
    # an input string and
    # spit RAID_info, V-disk 
    # info, and Time
    ##########################
    def extract_raid_vdisk_time(input_string):
        
        ############################
        # Extract Vdisk et al info
        ############################

        temp_index = input_string.find('(', 0, len(input_string))

        temp_c1 = input_string[temp_index:].strip('\n')
        temp_index_c1 = temp_c1.find(', ', 0, len(temp_c1))

        temp_c2 = temp_c1[1:temp_index_c1]
        temp_index_c2 = temp_c1.find(':', 0, len(temp_c1))

        vdisk_info = temp_c2[temp_index_c2+1:]

        
        ############################
        # Extract RAID et al info
        ############################

        temp_index_c2 = temp_c1.find(')', 0, len(temp_c1))

        raid_info = temp_c1[temp_index_c2+3:-1]
        
        
        ############################
        # Extract Date & Time
        ############################

        temp_index1 = input_string.find(' ', 0, len(input_string))

        temp_c2 = (input_string[temp_index1:].strip(' '))

        temp_index1 = temp_c2.find('  ', 0, len(temp_c2))

        
        import time

        string_time = time.strptime(str(temp_c2[:temp_index1]),
                                        "%Y-%m-%d %H:%M:%S")
         
        # convert it in seconds
        string_time = time.mktime(string_time) 
        
        return [vdisk_info, raid_info, string_time]    
        
        
        
    #########################
    # It will calculate
    # Build time from creation
    # and completion time
    #########################
    def calculate_build_time(creation_time, completion_time):
        
        temp = completion_time - creation_time
        
        hours = int(int(temp) / 3600 % 24)
        minutes = int(int(temp) / 60 % 60)
        seconds = int(int(temp) % 60)
        days = int(int(temp) / 86400)
        
        build_time_string = str(days) + ':' + str(hours) + ':' + str(minutes) + ':' +str(seconds) 
                             
        return str(build_time_string)  

    
    
    #######################
    # It will check if all
    # disks are created, if
    # all disks are created 
    # or reconstructed. It 
    # will notify User.
    #######################
    def check_len(creation_start, creation_complete, in_string):                                        
        ##################################
        # To check if number of creation
        # is same number of completion
        ##################################
        flag = 0
        if len(creation_start) == len(creation_complete):
            pass
        else:
            # Raise flag if lengths are not same 
            flag = 1
            print("Some Disk's " +str(in_string)+" is incomplete!" + ' Start: '+str(len(creation_start))+', Complete: '+str(len(creation_complete)))        
        
        return flag
    

    
    ###########################
    # Check for correct
    # V-disk number and serial 
    # number to calculate 
    # Build Time
    ###########################
    def extract_final_data(creation_start, creation_complete):
    
        vdisk_raid_buildtime_info = []

        for i in range(len(creation_start)):

            # Extract start info
            [vdisk_start,
            raid_start, 
            start_time] = Raid_Fault_Functions.extract_raid_vdisk_time(creation_start[i])
            
           
            for j in range(len(creation_complete)):

                # Extract start info
                [vdisk_complete,
                 raid_complete, 
                 complete_time] = Raid_Fault_Functions.extract_raid_vdisk_time(creation_complete[j])
                
                if vdisk_complete == vdisk_start :

                    build_time = Raid_Fault_Functions.calculate_build_time(start_time, complete_time)

                    #vdisk_raid_buildtime_info.append(vdisk_start)
                    #vdisk_raid_buildtime_info.append(raid_start)
                    vdisk_raid_buildtime_info.append([vdisk_start, build_time])
                    #vdisk_raid_buildtime_info.append('\n')

                    break
                    
        return vdisk_raid_buildtime_info  



    
    ############################################
    # This function will give  
    # Creation & Reconstruction information
    # of given controller data.
    # 
    # 1.) Vdisk information et. al
    # 2.) RAID info et. al 
    # 3.) Build time(Time difference between
    #     Start and completion of build)
    ############################################
    def get_info_of_controller(c_path, 
                               controller_data, 
                               creation_start_string, 
                               creation_complete_string,
                               recreation_start_string,
                               recreation_complete_string):
                               
        [creation_start, creation_complete, 
         recreation_start, recreation_complete] = Raid_Fault_Functions.extract_create_recreation_data (c_path, 
                                                controller_data, creation_start_string, 
                                                creation_complete_string,
                                                recreation_start_string,
                                                recreation_complete_string)

        ######################
        # CREATION                                        
        ######################
        repeatation_flag = Raid_Fault_Functions.check_len(creation_start, creation_complete, 'creation')
        
        if repeatation_flag == 1:
            creation_start = Raid_Fault_Functions.eliminate_repeatation(creation_start)               
            creation_complete = Raid_Fault_Functions.eliminate_repeatation(creation_complete)               

        
        final_creation_data = Raid_Fault_Functions.extract_final_data(creation_start, creation_complete)
        #print('\n\n' + creation_start[0] +'\n\n Complete:'+creation_complete[0])
        
        ################################
        # RECONSTRUCTION / RECREATION                                        
        ################################
        repeatation_flag = Raid_Fault_Functions.check_len(recreation_start, recreation_complete, 'reconstruction')

        if repeatation_flag == 1:
            recreation_start = Raid_Fault_Functions.eliminate_repeatation(recreation_start)               
            recreation_complete = Raid_Fault_Functions.eliminate_repeatation(recreation_complete)
        
        final_recreation_data = Raid_Fault_Functions.extract_final_data(recreation_start, recreation_complete)
                        
                       
                        
        return [final_creation_data, final_recreation_data]
        
    
    
    ################################
    # It will extract controller 
    # information from log file(s)
    ################################
    def extract_controller_name(data, indices):
        
        controller_name = [data[indices[i]] for i in range(len(indices))]
        
        log_controller_name = []
        for i in range(len(controller_name)):

            log_controller_name_temp = controller_name[i]

            temp_index = log_controller_name_temp.find('--', 0, len(log_controller_name_temp))

            log_controller_name.append(log_controller_name_temp[temp_index+3:].strip('\n'))
        return log_controller_name    
        
        
        
    ################################
    # Extract vdisk, raid, and 
    # build time for both controllers
    ################################    
    def generate_vdisk_raid_build_info(lf, c_path, log_data):
        
        import sys
        
        #################################
        # Qualifying Controller String
        # to reduce the search space
        #################################

        # Finding index where this string
        controller_string ='SC Event Log -- Controller '

        controller_index = lf.find_index(log_data,controller_string) 
        #print(controller_index)
        
        log_controller_name = Raid_Fault_Functions.extract_controller_name(log_data, controller_index)
        
        no_of_controller = len(controller_index)

        # String which indicates creation started
        creation_start_string = 'Vdisk creation started.'
        # String which indicates creation completed successfully
        creation_complete_string = 'Vdisk creation completed successfully.'

        # String which indicates recreation started
        recreation_start_string = 'Vdisk reconstruction started.'
        # String which indicates recreation completed successfully
        recreation_complete_string = 'Reconstruction of a vdisk completed.'

        if no_of_controller == 0: # 0 controller
            print('\n\n No controllers found!\
                   \n\n\nPROCESS TERMINATED !\
                   \n\nPlease check if "SC Event Log -- Controller "\
                   string is present.')
                   
            sys.exit()  
            
        elif no_of_controller == 1: # 1 controller
            print('\n\nOnly one controller present.')

            # 1st controller search space
            controller_1_ss = log_data[controller_index[0] : controller_index[1]]
            
            controller_2_ss = []
            
            #print('c1')
            [c1_creation_data, c1_recreation_data] = Raid_Fault_Functions.get_info_of_controller(c_path, 
                                                                               controller_1_ss, 
                                                                               creation_start_string, 
                                                                               creation_complete_string,
                                                                               recreation_start_string,
                                                                               recreation_complete_string)                
            
            c2_creation_data = [] 
            c2_recreation_data = []

        elif no_of_controller == 2:  # 2 controllers
            
            # 1st controller search space
            controller_1_ss = log_data[controller_index[0] : controller_index[1]]

            # 2nd controller search space
            controller_2_ss = log_data[controller_index[1]: ]
            
            #print('c1')
            [c1_creation_data, c1_recreation_data] = Raid_Fault_Functions.get_info_of_controller(c_path, 
                                                                               controller_1_ss, 
                                                                               creation_start_string, 
                                                                               creation_complete_string,
                                                                               recreation_start_string,
                                                                               recreation_complete_string)                
            #print('c2')
            [c2_creation_data, c2_recreation_data] = Raid_Fault_Functions.get_info_of_controller(c_path, 
                                                                               controller_2_ss, 
                                                                               creation_start_string, 
                                                                               creation_complete_string,
                                                                               recreation_start_string,
                                                                               recreation_complete_string)                


        else: # more than 2 controllers
            print('\n\nMore than 2 controllers found!\
                   \n\n\nPROCESS TERMINATED !')
                   
            sys.exit()       


        return [ no_of_controller, log_controller_name,
                c1_creation_data, c1_recreation_data, 
                c2_creation_data, c2_recreation_data]
        
        
        
    #################################
    # This function will eliminate
    # any repeatation
    #################################
    def eliminate_repeatation(data_list):
        
        # Extract vdisk number to check repeatation
        vdisk_tem = []
        for i in range(len(data_list)):

            vdisk_tem.append(Raid_Fault_Functions.extract_raid_vdisk_time(data_list[i])[0])
            
            
        # collect non repeated elements in nonrepeat_list
        nonrepeat_list = []
        for i in vdisk_tem:
            if i not in nonrepeat_list:
                nonrepeat_list.append(i)
        
        
        # collect Indices of each repeated elements
        repeatation_index = []
        for i in range(len(nonrepeat_list)):
            repeatation_index_temp = []
            
            for j in range(len(data_list)):

                if nonrepeat_list[i] in data_list[j]:
                    repeatation_index_temp.append(j)
                    #break
            repeatation_index.append(repeatation_index_temp)

        # if more each element is repeated, consider only first element
        new_replist = []
        for i in range(len(repeatation_index)):
            if len(repeatation_index[i]) > 1:
                temp = repeatation_index[i]
                repeatation_index[i] = [temp[0]] 
        
        new_data_list = [data_list[repeatation_index[i][0]] for i in range(len(repeatation_index))]
        
        return new_data_list

        
    ############################
    # This will simply write 
    # a list of strings in 
    # given document.
    ############################
    def write_list(document, data_list):
    
        from docx import Document
        from docx.shared import Inches
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.table import WD_TABLE_DIRECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.section import WD_SECTION
        from docx.shared import Cm
        import re
        
        for i in range(len(data_list)):
            temp_para= document.add_paragraph()
            temp_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = temp_para.add_run(data_list[i])
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(11)
        return document
        
        
        
    ############################
    # This will simply write 
    # a list of strings in 
    # given document with given 
    # Identation in inches.
    ############################    
    def write_list_with_indent(document, data_list, num):

        from docx import Document
        from docx.shared import Inches
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.table import WD_TABLE_DIRECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.section import WD_SECTION
        from docx.shared import Cm
        import re
        
    
        # temp_para=document.add_paragraph()
        
        # paragraph_format = temp_para.paragraph_format
        # paragraph_format.left_indent
        # paragraph_format.left_indent = Inches(num)
        
        for i in range(len(data_list)):
            temp_para= document.add_paragraph()
            paragraph_format = temp_para.paragraph_format
            paragraph_format.left_indent
            paragraph_format.left_indent = Inches(num)

            run = temp_para.add_run(data_list[i])
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(11)
        return document
        
        
        
    ############################
    # This function will extract
    # non-repeated list of 
    # vdisk's status.
    ############################
    def extract_vdisk_status(lf, log_data):
        
        ##############################
        #  Collecting vdisk critical
        #  data.
        #############################
        critical_string1 = 'CRITICAL       A vdisk is '
        critical_indices1 = lf.find_index(log_data, critical_string1)
        critical_data1 = [log_data[critical_indices1[i]] for i in range(len(critical_indices1))]

        critical_string2 = 'WARNING       A vdisk is '
        critical_indices2 = lf.find_index(log_data, critical_string2)
        critical_data2 = [log_data[critical_indices2[i]] for i in range(len(critical_indices2))]

        critical_data = critical_data1 + critical_data2
        

        ########################
        # Grabing vdisk status 
        # which says: Critical
        # or Degraded.
        ########################
        vdisk_stat = []
        for i in range(len(critical_data)):

            ################################
            # String "A vdisk is critical" or
            # "A vdisk is degraded". So just 
            # to grab degraded or critical
            # find index of "A vdisk is ", which
            # is 11 characters. And index of "."
            # (period). 
            ################################
            temp_cd = critical_data[i]

            temp_index1 = temp_cd.find('A vdisk is ', 0, len(temp_cd))
            temp_index2 = temp_cd.find('.', 0, len(temp_cd))
            vdisk_status = temp_cd[temp_index1+11:temp_index2]
            vdisk_status = vdisk_status.upper()
            
            # v-disk number
            temp_index3 = temp_cd.find('vdisk: ', 0, len(temp_cd))
            temp_index4 = temp_cd.find(',', 0, len(temp_cd))
            vdisk_no = temp_cd[temp_index3+7:temp_index4]

            vdisk_stat.append([vdisk_no, vdisk_status])
        
        
        #####################
        # Eliminating repeated 
        # values
        #####################
        new_vdisk_status = []
        for i in vdisk_stat:
            if i not in new_vdisk_status:
                new_vdisk_status.append(i)
        
        return new_vdisk_status
        
        
        
    ######################
    # TABLE - 1
    # VDISK, LOCATION, &
    # SERIAL NUMBER DATA
    ######################
    def extract_t1_vdisk_location_serialno(lf, log_data):
        
        ##################
        # Grabbing the 
        # 1st Table
        ##################
        ua = log_data
        dpsld_1st_str ='Location'
        end_str= '------------------------------------------------------------------------------------\n'

        sw_index=[]

        for j in range(1,len(ua)):
            if dpsld_1st_str in str(ua[j]) :
                sw_index.append(j)
                #print(ua[j])

        #print(len(sw_index))
        serial_str='Serial Number'
        v_str='Vendor'
        rev_str='Rev'
        des_str= 'Description'
        usg_str= 'Usage'

        #print(sw_index)
        sw_index1=[]
        for i in range(len(sw_index)):
            temp_data = ua[sw_index[i]]
            if serial_str and v_str and rev_str in temp_data: 
                #if  in temp_data: 
                    #if  in temp_data:
                        sw_index1.append(i)
        #print(sw_index1)                    
        new_dpsld_1st =ua[sw_index[sw_index1[0]]:]



        end_str_index_1st=lf.find_index(new_dpsld_1st
                                                    ,end_str)

        dpsld_1st_data= new_dpsld_1st[end_str_index_1st[0]-1:
                                        end_str_index_1st[1]]

        dpsld_1st_data= lf.strip_extras(dpsld_1st_data)

        ################################
        # To save row index which has 
        # Data present (i.e Eliminating 
        # Not present/available data)
        ################################
        req_row_index=[] 
        for i in range(len(dpsld_1st_data)):
            if 'Not Present' not in dpsld_1st_data[i]:
                    req_row_index.append(i)

                    
        dpsld_1st_data_new = [dpsld_1st_data[req_row_index[i]] for i in range(len(req_row_index))]            
        
        dpsld_1st_data = dpsld_1st_data_new
        
        
        ##############################
        # In table 1 after "##", lines
        # does not have data, this code
        # will eliminate it, if first
        # location is not a number
        ##############################

        final_data = []
        for i in range(len(dpsld_1st_data)): 
            single_line = dpsld_1st_data[i]

            sl_index = single_line.find(',', 0 , len(single_line))
            #print(single_line[ :sl_index])
            try:
                if ((str(int(float(single_line[ :sl_index]))))).isnumeric():
                    final_data.append(single_line)
            except:
                continue
        final_data = [dpsld_1st_data[0]] + final_data
        
        
        
        #########################
        # Extract VDISK, LOCATION 
        # , and SERIAL NO. DATA
        #########################
        t1_vdisk_location_serialno = []
        for i in range(1, len(final_data)):

            # For location
            location_index = final_data[i].find(',', 0, len(final_data[i]))

            # For Vdisk
            # begin index
            vdisk_index_b = final_data[i].find('vdisk-', 0, len(final_data[i]))
            vdisk_temp = (final_data[i])[vdisk_index_b:]
            #print(vdisk_temp)

            # stop index
            vdisk_index_s = vdisk_temp.find(',', 0, len(vdisk_temp))
            #vdisk_data.append(vdisk_temp[:vdisk_index_s])

            vdisk_data = vdisk_temp[:vdisk_index_s]
            if vdisk_data == '':
                vdisk_data='Available'

            # Serial Number
            # begin index is location index
            serial_temp = (final_data[i])[location_index+1:]
            #print(serial_temp)
            # stop index
            serial_index_s = serial_temp.find(',', 0, len(serial_temp))


            t1_vdisk_location_serialno.append([vdisk_data, (final_data[i])[:location_index], serial_temp[:serial_index_s]])

        #location_data, vdisk_data, 
        return t1_vdisk_location_serialno

    
    
    ##################################
    # TABLE - 2
    # Pulling out Model No. and Size
    # associated with vdisk number.
    ##################################
    def extract_t2_vdisk_modelno_size(lf, log_data, t1_vdisk_location_serialno):
        
        ua = log_data
        dpsld_1st_str ='Status'
        end_str= '------------------------------------------------------------------------------------\n'

        sw_index=[]
        #print(len(ua))

        for j in range(1,len(ua)):
            if dpsld_1st_str in str(ua[j]) :
                sw_index.append(j)
                #print(ua[j])

        #print(len(sw_index))
        encl_str='Encl'
        slot_str='Slot'
        v_str='Vendor'
        #raid_str= 'RAID'
        model_str= 'Model'
        sn_str= 'Serial Number'

        #print(sw_index)
        sw_index1=[]
        for i in range(len(sw_index)):
            temp_data = ua[sw_index[i]]
            if encl_str and slot_str and v_str and model_str and sn_str in temp_data: 
                #if  in temp_data: 
                    #if  in temp_data:
                        sw_index1.append(i)
        #print(sw_index1)                    
        new_dpsld_1st =ua[sw_index[sw_index1[0]]:]

        
        #new_dpsld_1st = ua[sw_index[0]:]

        end_str_index_1st = lf.find_index(new_dpsld_1st, end_str)

        dpsld_1st_data= new_dpsld_1st[end_str_index_1st[0]-1:
                                        end_str_index_1st[1]]
        
        ################################
        # To save row index which has 
        # Data present (i.e Eliminating 
        # Not present/available data)
        ################################
        req_row_index=[] 
        for i in range(len(dpsld_1st_data)):
            if 'Not Present' not in dpsld_1st_data[i]:
                    req_row_index.append(i)

                    
        dpsld_1st_data_new = [dpsld_1st_data[req_row_index[i]] for i in range(len(req_row_index))]            
        
        dpsld_1st_data = dpsld_1st_data_new
        ########################
        # Finding index of MOdel
        # and Size
        ########################
        model_index = dpsld_1st_data[0].find('Model', 0, len(dpsld_1st_data[0]))
        serial_index = dpsld_1st_data[0].find('Serial', 0, len(dpsld_1st_data[0]))
        size_index = dpsld_1st_data[0].find('Size', 0, len(dpsld_1st_data[0]))

        
        t2_vdisk_model_size = []
        for i in range(2, len(dpsld_1st_data)):

            # Model Number
            model_data = (dpsld_1st_data[i])[model_index: serial_index]
            model_data = model_data.strip(' ')

            # Size / Capacity
            size_data = (dpsld_1st_data[i])[size_index:]
            size_data = size_data.strip('\n')

            # Serial Number
            serial_data = (dpsld_1st_data[i])[serial_index: size_index]
            serial_data = serial_data.strip(' ')
            
            for k in range(len(t1_vdisk_location_serialno)):
                if serial_data == (t1_vdisk_location_serialno[k])[2]:
                    t2_vdisk_model_size.append([(t1_vdisk_location_serialno[k])[0], model_data, size_data])
                    break
            
        return t2_vdisk_model_size
    
    
    
    
    ############################
    # TABLE - 3
    # Pulling RAID, Own, and 
    # Status
    ############################
    def extract_t3_vdisk_raid_own_stat(lf, log_data):
        
        ua = log_data
        dpsld_1st_str ='Name'
        end_str= '------------------------------------------------------------------------------------\n'

        sw_index=[]
        #print(len(ua))

        for j in range(1,len(ua)):
            if dpsld_1st_str in str(ua[j]) :
                sw_index.append(j)
                #print(ua[j])
        
        #print(len(sw_index))
        size_str='Size'
        free_str='Free'
        own_str='Own'
        raid_str= 'RAID'
        status_str= 'Status'

        #print(sw_index)
        sw_index1=[]
        for i in range(len(sw_index)):
            temp_data = ua[sw_index[i]]
            if size_str and free_str and own_str and raid_str and status_str in temp_data: 
                #if  in temp_data: 
                    #if  in temp_data:
                        sw_index1.append(i)
        #print(sw_index1)                    
        new_dpsld_1st =ua[sw_index[sw_index1[0]]:]


        #new_dpsld_1st =ua[sw_index[-1]:]

        end_str_index_1st=lf.find_index(new_dpsld_1st, end_str)
                                                    

        dpsld_1st_data= new_dpsld_1st[end_str_index_1st[0]-1:
                                        end_str_index_1st[1]]

        
        ########################
        # Finding index of MOdel
        # and Size
        ########################
        # to extract vdisk info
        name_index = dpsld_1st_data[0].find('Name', 0, len(dpsld_1st_data[0]))
        size_index = dpsld_1st_data[0].find('Size', 0, len(dpsld_1st_data[0]))

        # to extract Ownership info
        own_index = dpsld_1st_data[0].find('Own', 0, len(dpsld_1st_data[0]))
        pref_index = dpsld_1st_data[0].find('Pref', 0, len(dpsld_1st_data[0]))

        # to extract RAID info
        raid_index = dpsld_1st_data[0].find('RAID', 0, len(dpsld_1st_data[0]))
        class_index = dpsld_1st_data[0].find('Class', 0, len(dpsld_1st_data[0]))


        # to extract vdisk status info
        status_index = dpsld_1st_data[0].find('Status', 0, len(dpsld_1st_data[0]))
        job_index = dpsld_1st_data[0].find('Job', 0, len(dpsld_1st_data[0]))

        t3_vdisk_raid_own_status = []
        for i in range(2, len(dpsld_1st_data)):

            # vdisk data
            vdisk_data = (dpsld_1st_data[i])[name_index: size_index]
            vdisk_data = vdisk_data.strip(' ')

            # ownership data
            own_data = (dpsld_1st_data[i])[own_index: pref_index]
            own_data = own_data.strip(' ')

            # raid data
            raid_data = (dpsld_1st_data[i])[raid_index: class_index]
            raid_data = raid_data.strip(' ')

            # status data
            status_data = (dpsld_1st_data[i])[status_index: job_index]
            status_data = status_data.strip(' ')

            t3_vdisk_raid_own_status.append([vdisk_data, raid_data, own_data, status_data])

        return t3_vdisk_raid_own_status
    
    
    
    ############################
    # Create and write .csv
    # in required format.
    #
    # It will then used to 
    # populate table in report.
    ############################
    def generate_time_csv(lf, log_data, part_no, script_dir, c1_creation_data, c1_recreation_data, c2_creation_data, c2_recreation_data):
        
        # Merge Controller 1 & 2 data
        creation_data = c1_creation_data + c2_creation_data
        recreation_data = c1_recreation_data + c2_recreation_data
        
        t1_data = Raid_Fault_Functions.extract_t1_vdisk_location_serialno(lf, log_data)
        t2_data = Raid_Fault_Functions.extract_t2_vdisk_modelno_size(lf, log_data, t1_data)    
        t3_data = Raid_Fault_Functions.extract_t3_vdisk_raid_own_stat(lf, log_data)
        t4_data = Raid_Fault_Functions.extract_vdisk_status(lf, log_data)
        
 
        #########################
        # Writing .csv file
        # in required format
        #########################
        temp_name = '\\' +str(part_no)+ '_temp'
        t4_vdisk_list = [t4_data[k][0] for k in range(len(t4_data))]

        t5_creation_list = [creation_data[k][0] for k in range(len(creation_data))]
        t5_recreation_list = [recreation_data[k][0] for k in range(len(recreation_data))]
        
        
        
        with open ( r''+str(script_dir)+ str(temp_name)+ '.csv', "w" ) as out_file:

            out_string = ""
            out_string += "Vdisk #,RAID #,Model #,Capacity,Controller,Location,Construction Time (D:H:M:S),Status,Reconstruction Time (D:H:M:S),Status\n"
            
            for i in range(len(t1_data)):

                # vdisk number
                out_string += t1_data[i][0]
                for j in range(len(t3_data)):


                    if t1_data[i][0] == t3_data[j][0]:

                        # RAID Number
                        out_string += ',' + t3_data[j][1]

                        # Model Number
                        out_string += ',' + t2_data[i][1]

                        #  Size
                        out_string += ',' + t2_data[i][2]
                        #out_string += ',' + t1_data[i][1]

                        # Ownership; which controller 
                        #owns this disk
                        out_string += ',' + t3_data[j][2]

                        # Location
                        out_string += ',' + t1_data[i][1]

                        # Creation time
                        t5_creation_list
                        if t1_data[i][0] in t5_creation_list:
                            ct_index = t5_creation_list.index(t1_data[i][0])
                            out_string += ',' + creation_data[ct_index][1]
                        else:
                            out_string += ',N.A.' 

                        # Status (Critical/Degraded)
                        #for k in range(len(t4_data)):
                        if t1_data[i][0] in t4_vdisk_list:
                            vd_index = t4_vdisk_list.index(t1_data[i][0])
                            out_string += ',' + t4_data[vd_index][1]
                        else:
                            out_string += ',N.A.' 

                        # Recreation time
                        t5_recreation_list
                        if t1_data[i][0] in t5_recreation_list:
                            rt_index = t5_recreation_list.index(t1_data[i][0])
                            out_string += ',' + recreation_data[rt_index][1]
                        else:
                            out_string += ',N.A.' 

                        # Status (i.e. FTOL)
                        out_string += ',' + t3_data[j][3]

                        out_string += '\n'



                    elif t1_data[i][0] == 'Available':
                        # RAID Number
                        out_string += ',N.A.' 
                        #+ t3_data[j][1]

                        # Model Number
                        out_string += ',' + t2_data[i][1]

                        #  Size
                        out_string += ',' + t2_data[i][2]
                        #out_string += ',' + t1_data[i][1]

                        # Ownership; which controller 
                        #owns this disk
                        out_string += ',N.A.' 
                        #+ t3_data[j][2]

                        # Location
                        out_string += ',' + t1_data[i][1]

                        out_string += ',N.A.,N.A.,N.A.,N.A.\n '
                        # Status (i.e. FTOL)
                        #out_string += ',' + t3_data[j][3]

                        #out_string += '\n'
                        break


            out_file.write(out_string)
    
    
#####################################
#              END                  #
#####################################    