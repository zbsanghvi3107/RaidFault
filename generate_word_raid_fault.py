####################################################
#                  Revision: 1.0                   #
#              Updated on: 11/12/2015              #
####################################################

####################################################
#                                                  #
#   This file writes/append data to the Template   #
#   Word File.                                     #
#                                                  #
#   Author: Zankar Sanghavi                        #
#                                                  #
#   © Dot Hill Systems Corporation                 #
#                                                  #
####################################################

import os
import sys

import raid_fault_functions
rff = raid_fault_functions.Raid_Fault_Functions

###################################
#  Importing from other Directory
###################################
import sys
os.chdir('..')
c_path = os.getcwd()

sys.path.insert(0, r''+str(c_path)+'/Common Scripts')
import user_inputs_ICS
import fixed_data_ICS
import report_functions
import extract_lists
import modify_word_docx

sys.path.insert(0, r''+str(c_path)+'/IO Stress')
import log_functions


###################################
#  Importing from Current Directory
###################################
sys.path.insert(0, r''+str(c_path)+'/Cable Pulls')
import pandas
import csv
import numpy as np

import extract_f2_f3
import time
import zipfile

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.shared import Cm
import re


###################################
#  To call functions from different
#  files. 
###################################
lf=log_functions.Log_Functions
ed= extract_f2_f3.Extract_F2_F3
ui= user_inputs_ICS.User_Inputs
fd= fixed_data_ICS.Fixed_Data
rf= report_functions.Report_Functions

class Generate_Word_Raid_Fault:

    ####################################################
    #   This Function does:
    #   
    #   1.) Uses "unzip_pull_log" function to unzip and
    #   find the .logs file in the folder to read the 
    #   data.
    #
    #   2.) Checks Errors, extracts: Errors info, Construction 
    #   & Reconstruction time Hardware, and Software. 
    #   
    #   3.) It writes required data, with formatting
    #   into a WORD report. 
    #   
    ####################################################
    def generate_final_report(rff, fixed_dir
                            , part_no, test_name, no_of_files
                            , csv_file_list, log_list
                            , fw_no_names, chassis_names
                            , cntrllr_names, t_fw_type):
        
        document = Document(r''+str(fixed_dir)+'\\'
                            +str(part_no)+str(test_name)
                            +'.docx')
                            
        document.add_page_break()
        
        f=no_of_files
        
        error_collection=[]
        
        for f in range(no_of_files):
            progress=(round((float(100/int(no_of_files))*f),2))
            # to show Report Progress
            
            print('\nReport Progess: ',progress,'%\n')
            
            log_temp = str(no_of_files)+' - Log file'
            log_data= lf.unzip_pull_log(log_list[f], 'store', log_temp) 
            #pull .logs file from zip folder
            
            
            ###################################
            #  Error check and data extraction
            ###################################
            [write_sum, read_sum, iter_flag, hw_list
            , host_list, sasmap_list]= ed.generate_f2_f3_iter(
                                       csv_file_list[f], log_data)

            
            ###################
            #   Summary
            ###################
            #document.add_page_break()
            
            document.add_heading('Raid Fault test Summary for '
                                + str(fw_no_names[f])  
                                + '\\' + str(chassis_names[f])
                                +  '\\'  + str(cntrllr_names[f])
                                + ' chassis' ,level=3)
            
            temp_para=document.add_paragraph()
            
            paragraph_format = temp_para.paragraph_format
            paragraph_format.left_indent
            paragraph_format.left_indent = Inches(0.5)
            
            run = temp_para.add_run('Read error(s): '
                                    +str(read_sum)
                                    + '\nWrite error(s): '
                                    + str(write_sum)
                                    +'\n')
                                    
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(11)
            
            
            ######################
            # Writing table
            ######################
            
            [ no_of_controller, log_controller_name,
                c1_creation_data, c1_recreation_data, 
                c2_creation_data, c2_recreation_data] = rff.generate_vdisk_raid_build_info(lf, c_path, log_data)
                
            # arrange all data to write in a table.
            rff.generate_time_csv(lf, log_data, part_no, fixed_dir, c1_creation_data, c1_recreation_data, c2_creation_data, c2_recreation_data)    
            
            

            document.add_heading('Disk Construction and Reconstruction details for '+ str(fw_no_names[f])  
                                +'\\' + str(chassis_names[f])
                                +  '\\'  + str(cntrllr_names[f])
                                + ' chassis', level=3)
                  
            
            file_name1 = '\\' +str(part_no)+ '_temp'
            file_data = pandas.read_csv(open(r''+ str(fixed_dir) 
                                        +str(file_name1) +'.csv')
                                        ,header=None)
                                        
            file_data = np.array(file_data)

            
            no_rows= len(file_data[:,0]) #finding number of rows
            no_columns = len(file_data[0,:]) #finding number of columns
            
            
            ###########################
            #   Creating a table and 
            #   setting its attributes.
            ###########################
            table = document.add_table(rows = no_rows, cols = no_columns)
                
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.table_direction = WD_TABLE_DIRECTION.LTR
            table_style = document.styles["Normal"]
            table_font = table_style.font
            table.style = 'Table Grid'
            table.autofit = True
            
                        
            hdr_cells=table.rows[0].cells
            
            # Populating a table from .csv file data. 
            for i in range(no_rows):
                for j in range(no_columns):
                    
                    
                    table_font.size = Pt(9)
                    hdr_cells = table.rows[i].cells
                    #hdr_cells[j].text.bold = True
                    hdr_cells[j].text = str(file_data[i,j])
                        
            
            if os.path.isfile(r''+ str(fixed_dir) +str(file_name1) +'.csv'):
                    
                os.remove(r''+ str(fixed_dir) +str(file_name1) +'.csv')
            
            # if no_of_controller == 1:
                    
                # document.add_heading('Disk creation details - '+str(log_controller_name[0])+ ': '
                                     # ,level=4)
                            
                # document = rff.write_list_with_indent(document, c1_creation_data, 0.6)
        
                
                # document.add_heading('Disk reconstrunction details - '+str(log_controller_name[0])+ ': '
                                     # ,level=4)
                            
                # document = rff.write_list_with_indent(document, c1_recreation_data, 0.6)
        
            # elif no_of_controller == 2:
               
                # #rff.generate_time_csv(log_data, script_dir, c1_creation_data, c1_recreation_data, c2_creation_data, c2_recreation_data)    
                
                
                # # Controller - 1
                # document.add_heading('Disk creation details - '+str(log_controller_name[0])+ ': '
                                     # ,level=4)
                            
                # document = rff.write_list_with_indent(document, c1_creation_data, 0.6)
        
                
                # document.add_heading('Disk reconstruction details - '+str(log_controller_name[0])+': '
                                     # ,level=4)
                            
                # document = rff.write_list_with_indent(document, c1_recreation_data, 0.6)
        
                
                # Controller - 2
                # document.add_heading('Disk creation details - '+str(log_controller_name[1])+': '
                                     # ,level=4)
                            
                # document = rff.write_list_with_indent(document, c2_creation_data, 0.6)
        
                
                # document.add_heading('Disk reconstruction details - '+str(log_controller_name[1])+': '
                                     # ,level=4)
                            
                # document = rff.write_list_with_indent(document, c2_recreation_data, 0.6)
                
                
            ###################
            #   F2 Menu
            ###################
            document.add_page_break()
            document.add_heading('Raid Fault (F2 menu) for '
                                + str(fw_no_names[f])  
                                +'\\' + str(chassis_names[f])
                                +  '\\'  + str(cntrllr_names[f])
                                + ' chassis' ,level=3)
            
            document = rff.write_list(document, hw_list)
                                         
            
            ###################
            #   F3 Menu
            ###################
            document.add_page_break()
            
            document.add_heading('Raid Fault (F3 menu) for '
                                + str(fw_no_names[f]) 
                                + '\\' + str(chassis_names[f])
                                +  '\\'  + str(cntrllr_names[f]) 
                                + ' chassis',level=3)
            
            document = rff.write_list(document, sasmap_list)
            
            # save the document
            document.save(r''+str(fixed_dir)+'\\'
                         +str(part_no)+str(test_name)+'.docx')
        
        return [ c1_creation_data, c1_recreation_data, 
                c2_creation_data, c2_recreation_data]        
        
#####################################
#              END                  #
#####################################